from docx import Document
from docx.shared import Pt

def write(name_doc, data, directory, name):
    document = Document(name_doc)
    doc = document
    paragraphs = doc.paragraphs
    print('paragraphs:', len(paragraphs))

    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)

    table = doc.tables[1]
    i = 2
    for el in data:
        print('el', el["name"])
        table.rows[i].cells[1].text = 'Статья'
        table.rows[i].cells[2].text = el["name"] + ". " + el["journal"]
        table.rows[i].cells[3].text = ", ".join(el["authors"])
        table.add_row()
        i += 1

    document.save(directory + name + '.docx')